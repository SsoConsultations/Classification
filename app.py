import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.preprocessing import StandardScaler, OneHotEncoder, LabelEncoder
from sklearn.model_selection import train_test_split, GridSearchCV
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier, GradientBoostingClassifier
from sklearn.svm import SVC
from sklearn.metrics import accuracy_score, precision_score, recall_score, f1_score, roc_auc_score, confusion_matrix, classification_report, roc_curve, auc, make_scorer
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
import warnings
import re
from PIL import Image

# Suppress warnings for cleaner app output
warnings.filterwarnings("ignore")

# --- Session State Initialization ---
if "analysis_completed" not in st.session_state:
    st.session_state.analysis_completed = False
if "file_uploader_key" not in st.session_state:
    st.session_state.file_uploader_key = 0
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
if "X_processed" not in st.session_state:
    st.session_state.X_processed = None
if "y_processed" not in st.session_state:
    st.session_state.y_processed = None
if "label_encoder" not in st.session_state:
    st.session_state.label_encoder = None
if "original_df_with_target" not in st.session_state:
    st.session_state.original_df_with_target = None
if "feature_names" not in st.session_state:
    st.session_state.feature_names = []
if "target_name" not in st.session_state:
    st.session_state.target_name = None
if "evaluation_results_with_best_params" not in st.session_state:
    st.session_state.evaluation_results_with_best_params = {}
if "recommended_model_info" not in st.session_state:
    st.session_state.recommended_model_info = None


# Page config
st.set_page_config(
    layout="wide",
    page_title="Supervised Learning: Classification App"
)

# --- Authentication Logic ---
def login_page():
    """Displays the login page for user authentication."""
    try:
        logo = Image.open("SsoLogo.jpg")
        st.image(logo, width=150)
    except FileNotFoundError:
        st.warning("SsoLogo.jpg not found. Please ensure it's in the repository.")

    st.title("🔒 Login to Classification Dashboard")
    st.markdown("Please enter your credentials to access the application.")

    username = st.text_input("Username")
    password = st.text_input("Password", type="password")

    if st.button("Login"):
        try:
            correct_username = st.secrets["username"]
            correct_password = st.secrets["password"]
        except KeyError:
            st.error("Secrets not configured! Please set 'username' and 'password' in Streamlit Cloud secrets.")
            return

        if username == correct_username and password == correct_password:
            st.session_state.authenticated = True
            st.success("Login successful! Redirecting...")
            st.rerun()
        else:
            st.error("Invalid username or password.")

    st.markdown("---")
    st.markdown("<p style='text-align: center; color: grey;'>© Copyright SSO Consultants</p>", unsafe_allow_html=True)


# --- Helper Functions ---

@st.cache_data
def detect_potential_id_columns(df, uniqueness_threshold=0.9):
    """
    Detects columns that might be ID columns based on naming conventions and uniqueness.
    """
    potential_ids = []
    for col in df.columns:
        if any(keyword in col.lower() for keyword in ['id', 'user_id', 'customer_id', 'client_id', 'record_id']):
            potential_ids.append(col)
            continue
        temp_col = df[col].copy()
        try:
            numeric_check = pd.to_numeric(temp_col, errors='coerce')
            if not numeric_check.isnull().all():
                if numeric_check.nunique() / len(df) > uniqueness_threshold:
                    potential_ids.append(col)
                    continue
        except Exception:
            if df[col].dtype == 'object' and df[col].nunique() / len(df) > uniqueness_threshold:
                potential_ids.append(col)
                continue
    return list(set(potential_ids))

@st.cache_data
def preprocess_data(df, target_column_name, numeric_features, categorical_features, missing_strategy):
    """
    Preprocesses the data for classification: handles missing values, encodes categorical
    features, scales numeric features, and encodes the target variable if necessary.
    Returns processed features (X) and target (y), along with the LabelEncoder if used,
    and the original DataFrame rows that were kept after missing value handling.
    """
    df_proc = df.copy()

    y = df_proc[target_column_name]
    X = df_proc.drop(columns=[target_column_name])

    X_selected = X[numeric_features + categorical_features].copy()

    if missing_strategy == "drop_rows":
        combined_df = pd.concat([X_selected, y], axis=1)
        combined_df.dropna(inplace=True)
        X_selected = combined_df[X_selected.columns]
        y = combined_df[target_column_name]
    else: # impute
        for col in numeric_features:
            if col in X_selected.columns:
                X_selected[col].fillna(X_selected[col].mean(), inplace=True)
        for col in categorical_features:
            if col in X_selected.columns:
                X_selected[col].fillna(X_selected[col].mode()[0], inplace=True)
        if y.isnull().any():
            y.fillna(y.mode()[0], inplace=True)

    original_df_retained_rows = df_proc.loc[X_selected.index].copy()

    label_encoder = None
    if y.dtype == 'object' or pd.api.types.is_categorical_dtype(y):
        label_encoder = LabelEncoder()
        y_encoded = label_encoder.fit_transform(y)
    else:
        y_encoded = y.values

    encoded_feature_names = []
    if categorical_features:
        encoder = OneHotEncoder(sparse_output=False, handle_unknown="ignore")
        enc_data = encoder.fit_transform(X_selected[categorical_features])
        enc_df = pd.DataFrame(enc_data, columns=encoder.get_feature_names_out(categorical_features), index=X_selected.index)
        X_selected = X_selected.drop(columns=categorical_features)
        X_selected = pd.concat([X_selected, enc_df], axis=1)
        encoded_feature_names = encoder.get_feature_names_out(categorical_features).tolist()

    features_to_scale = numeric_features + encoded_feature_names
    
    X_processed = X_selected.copy()

    if features_to_scale:
        scaler = StandardScaler()
        X_processed[features_to_scale] = scaler.fit_transform(X_processed[features_to_scale])
    
    final_feature_names = X_processed.columns.tolist()

    return X_processed, y_encoded, label_encoder, original_df_retained_rows, final_feature_names

def format_metric(value):
    return f'{value:.4f}' if isinstance(value, (int, float)) else "N/A"

def add_styled_paragraph(document, text_content):
    """
    Adds a paragraph to the document, parsing text for Markdown-like bold (**text**) syntax
    and applying actual bold formatting.
    """
    p = document.add_paragraph()
    parts = re.split(r'(\*\*.*?\*\*)', text_content)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)


# Updated create_report function for Classification with proper classification report table
def create_report(document, algorithm, params, metrics, data_preview_df, confusion_matrix_plot_bytes, roc_curve_plot_bytes, feature_importance_plot_bytes, classification_report_dict, target_name, feature_names, class_labels, recommended_model_info=None):
    """Generates a comprehensive Word document report for classification."""

    # --- Add logo to report ---
    try:
        logo = Image.open("SsoLogo.jpg")
        logo_stream = io.BytesIO()
        logo.save(logo_stream, format="PNG")
        logo_stream.seek(0)
        document.add_picture(logo_stream, width=Inches(1.5))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except FileNotFoundError:
        document.add_paragraph("Logo (SsoLogo.jpg) not found for report.")

    document.add_heading('ML Classification Analysis Report', level=1)
    document.add_paragraph(f"Report generated on: {pd.to_datetime('today').strftime('%Y-%m-%d %H:%M:%S')}")

    document.add_heading('1. Analysis Overview', level=2)
    document.add_paragraph(
        "This report details the supervised classification analysis performed using the Streamlit application. "
        "The goal is to predict the target variable based on the selected features."
    )
    add_styled_paragraph(document, f"Target Variable: **{target_name}**")
    document.add_paragraph(f"Classes: {', '.join(class_labels)}")

    # --- New section for Model Recommendation in report ---
    if recommended_model_info:
        document.add_heading('2. Model Recommendation', level=2)
        add_styled_paragraph(document, f"Based on the automated evaluation, the recommended model is: **{recommended_model_info['name']}**")
        document.add_paragraph(f"Best parameters found: {', '.join([f'{k}={v}' for k, v in recommended_model_info['best_params'].items()])}")
        document.add_paragraph(f"Achieved {recommended_model_info['metric']}: {format_metric(recommended_model_info['score'])}")
        document.add_paragraph("These parameters serve as a strong starting point for further fine-tuning.")


    document.add_heading('3. Classification Parameters', level=2)
    document.add_paragraph(f"Algorithm Used: {algorithm}")
    for param, value in params.items():
        if value is not None:
            document.add_paragraph(f"- {param.replace('_', ' ').title()}: {value}")

    document.add_heading('4. Model Performance Metrics', level=2)
    document.add_paragraph(f"Accuracy: {format_metric(metrics.get('accuracy'))}")
    document.add_paragraph(f"Precision: {format_metric(metrics.get('precision'))}")
    document.add_paragraph(f"Recall: {format_metric(metrics.get('recall'))}")
    document.add_paragraph(f"F1-Score: {format_metric(metrics.get('f1_score'))}")
    if not np.isnan(metrics.get('roc_auc', np.nan)):
        document.add_paragraph(f"ROC AUC Score: {format_metric(metrics.get('roc_auc'))}")
    document.add_paragraph(
        "These metrics evaluate the performance of the classification model. "
        "Accuracy is the proportion of correct predictions. Precision is the ability of the classifier not to label as positive a sample that is negative. "
        "Recall is the ability of the classifier to find all the positive samples. F1-score is the weighted average of Precision and Recall. "
        "ROC AUC measures the area under the Receiver Operating Characteristic curve, indicating the model's ability to distinguish between classes (primarily for binary classification)."
    )

    document.add_heading('5. Detailed Classification Report', level=2)
    document.add_paragraph("This report shows the main classification metrics per class.")
    
    report_df = pd.DataFrame(classification_report_dict).transpose()
    if 'accuracy' in report_df.index:
        report_df = report_df.drop(index='accuracy')

    if all(col in report_df.columns for col in ['precision', 'recall', 'f1-score', 'support']):
        report_df = report_df[['precision', 'recall', 'f1-score', 'support']]

    for key in ['macro avg', 'weighted avg']:
        if key in classification_report_dict and isinstance(classification_report_dict[key], dict):
            report_df.loc[key] = [
                format_metric(classification_report_dict[key].get('precision')),
                format_metric(classification_report_dict[key].get('recall')),
                format_metric(classification_report_dict[key].get('f1-score')),
                str(classification_report_dict[key].get('support', 'N/A'))
            ]
    
    table = document.add_table(rows=1, cols=report_df.shape[1] + 1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Class/Metric"
    for i, col in enumerate(report_df.columns):
        hdr_cells[i+1].text = col
        hdr_cells[i+1].paragraphs[0].runs[0].font.bold = True

    for idx, row in report_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        for i, val in enumerate(row):
            cell_text = format_metric(val) if isinstance(val, (float, np.float64)) else str(val)
            row_cells[i+1].text = cell_text
            if isinstance(val, (float, np.float64, int, np.int64)):
                row_cells[i+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                row_cells[i+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_heading('6. Data Preview (First 5 Rows with Predicted Target)', level=2)
    data_table = document.add_table(rows=1, cols=data_preview_df.shape[1])
    data_table.style = 'Table Grid'
    data_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells_data = data_table.rows[0].cells
    for i, col in enumerate(data_preview_df.columns):
        hdr_cells_data[i].text = col
        hdr_cells_data[i].paragraphs[0].runs[0].font.bold = True

    for index, row in data_preview_df.iterrows():
        row_cells_data = data_table.add_row().cells
        for i, val in enumerate(row):
            row_cells_data[i].text = str(val)
            try:
                float(val)
                row_cells_data[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            except ValueError:
                row_cells_data[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT


    document.add_heading('7. Model Visualizations', level=2)
    if confusion_matrix_plot_bytes:
        document.add_paragraph("Confusion Matrix: Visualizes the performance of a classification model on a set of test data, showing true vs. predicted counts for each class.")
        document.add_picture(io.BytesIO(confusion_matrix_plot_bytes), width=Inches(5))
    else:
        document.add_paragraph("Confusion Matrix plot could not be generated.")

    if roc_curve_plot_bytes:
        document.add_paragraph("ROC Curve: Illustrates the diagnostic ability of a binary classifier system as its discrimination threshold is varied. A curve closer to the top-left corner indicates better performance.")
        document.add_picture(io.BytesIO(roc_curve_plot_bytes), width=Inches(5))
    else:
        document.add_paragraph("ROC Curve plot could not be generated (e.g., not binary classification or probabilities not available).")

    if feature_importance_plot_bytes:
        document.add_paragraph("Feature Importance Plot: Shows the relative importance of each feature in predicting the target variable. Higher importance means the feature had a greater impact on the model's decisions.")
        document.add_picture(io.BytesIO(feature_importance_plot_bytes), width=Inches(6))
    else:
        document.add_paragraph("Feature Importance plot could not be generated (e.g., model does not support it or no features selected).")

    document.add_heading('8. Prescriptive Insights', level=2)
    document.add_paragraph(
        "Based on the model's performance and the identified feature importances, here are some potential actionable insights. "
        "These insights aim to translate model findings into practical recommendations for business strategy. "
    )
    
    if feature_names and metrics.get('feature_importances') is not None and len(feature_names) == len(metrics['feature_importances']):
        importance_df_report = pd.DataFrame({
            'Feature': feature_names,
            'Importance': metrics['feature_importances']
        }).sort_values(by='Importance', ascending=False)
        
        top_features = importance_df_report['Feature'].head(5).tolist()
        
        add_styled_paragraph(document, f"**Key Drivers Identified:** The model indicates that features such as **{', '.join(top_features)}** are among the most influential in predicting **{target_name}**. Further investigation into these areas could yield significant insights.")
        add_styled_paragraph(document, f"**Potential Actions:** Consider strategies that target or leverage these key features. For example, if 'MonthlyCharges' is highly important for churn prediction, analyzing pricing strategies or offering tailored plans might be effective.")
        add_styled_paragraph(document, f"**Further Exploration:** Delve deeper into the characteristics of each predicted class (e.g., 'Churn' vs. 'No Churn' customers) based on these important features to craft highly targeted interventions.")
    else:
        document.add_paragraph("Specific prescriptive insights based on feature importance could not be generated as feature importance data was unavailable.")

    document.add_page_break()


# --- Main Application Content ---
def main_app():
    """Main application logic for the classification dashboard."""
    st.sidebar.title(" ")
    try:
        logo = Image.open("SsoLogo.jpg")
        st.sidebar.image(logo, use_container_width=True)
    except FileNotFoundError:
        st.sidebar.warning("SsoLogo.jpg not found. Please ensure it's in the repository.")

    st.title("📊 Supervised Learning: Classification App")

    st.markdown("""
    Welcome! This app helps you build and evaluate classification models to predict a target variable
    based on your data attributes.
    """)

    if st.session_state.authenticated:
        if st.sidebar.button("Logout"):
            st.session_state.authenticated = False
            st.rerun()

    st.header("1️⃣ Upload Your Data")
    st.markdown("""
    Upload your data file in **CSV or Excel** format.
    """)
    uploaded_file = st.file_uploader("Upload a CSV or Excel file", type=["csv", "xlsx"], key=f"file_uploader_{st.session_state.file_uploader_key}")

    df = None
    if uploaded_file:
        try:
            if uploaded_file.name.endswith(".csv"):
                df = pd.read_csv(uploaded_file)
            else:
                df = pd.read_excel(uploaded_file)
            st.success(f"✅ Successfully loaded {uploaded_file.name}")
        except Exception as e:
            st.error(f"Error loading file: {e}")

    if df is not None:
        st.header("2️⃣ Data Overview & Feature Selection")

        st.markdown("Here are the first 5 rows of your uploaded data:")
        st.dataframe(df.head())

        all_columns = df.columns.tolist()

        potential_id_cols = detect_potential_id_columns(df)

        st.subheader("Exclude Columns")
        st.markdown("""
        Select any columns that should **not** be used in the analysis (e.g., ID columns, irrelevant text fields).
        """)
        excluded_columns = st.multiselect(
            "Columns to Exclude",
            all_columns,
            default=potential_id_cols
        )

        df_filtered = df.drop(columns=excluded_columns)
        available_columns_for_selection = df_filtered.columns.tolist()

        if not available_columns_for_selection:
            st.warning("No columns left after exclusion. Please adjust excluded columns.")
            st.stop()

        st.subheader("Select Target Variable & Features")
        st.markdown("""
        Choose the column you want the model to **predict** (your dependent variable).
        """)
        target_column = st.selectbox(
            "Target Variable",
            available_columns_for_selection
        )

        if target_column:
            st.session_state.target_name = target_column

            y_raw = df_filtered[target_column]
            X_raw = df_filtered.drop(columns=[target_column])

            st.markdown("---")
            st.markdown("""
            Now, select the **independent features** (predictors) that the model will use to make predictions.
            """)

            initial_numeric_features = X_raw.select_dtypes(include=np.number).columns.tolist()
            initial_categorical_features = X_raw.select_dtypes(include='object').columns.tolist()

            selected_numeric_features = st.multiselect(
                "Numeric Features",
                initial_numeric_features,
                default=initial_numeric_features
            )
            selected_categorical_features = st.multiselect(
                "Categorical Features",
                initial_categorical_features,
                default=initial_categorical_features
            )

            if not selected_numeric_features and not selected_categorical_features:
                st.warning("Please select at least one numeric or categorical feature for classification.")
                st.stop()

            features_to_preprocess = selected_numeric_features + selected_categorical_features
            if not features_to_preprocess:
                st.error("No features selected for preprocessing. Please select some features.")
                st.stop()

            st.subheader("Handle Missing Data")
            st.markdown("""
            Choose how to handle rows with missing values in your selected features or target variable:
            - **Drop Rows:** Remove any rows that have missing values.
            - **Impute:** Fill missing numeric values with the column mean and missing categories/target with the most common value.
            """)
            missing_strategy = st.selectbox(
                "Missing Data Handling",
                ("drop_rows", "impute"),
                format_func=lambda x: x.replace("_", " ").title()
            )

            st.header("3️⃣ Data Preprocessing")
            with st.spinner("Preprocessing your data..."):
                try:
                    st.session_state.X_processed, st.session_state.y_processed, st.session_state.label_encoder, st.session_state.original_df_with_target, st.session_state.feature_names = \
                        preprocess_data(df_filtered, target_column, selected_numeric_features, selected_categorical_features, missing_strategy)

                    st.success("✅ Preprocessing complete!")
                    st.write("Here is a preview of your processed features (scaled and one-hot encoded where applicable):")
                    st.dataframe(st.session_state.X_processed.head())
                    st.write(f"Processed target variable unique values: {np.unique(st.session_state.y_processed)}")

                except Exception as e:
                    st.error(f"Error during data preprocessing: {e}. Please check your data and selections.")
                    st.session_state.X_processed = None
                    st.session_state.y_processed = None
                    st.stop()

            if st.session_state.X_processed is not None and st.session_state.y_processed is not None:
                if st.session_state.X_processed.shape[0] < 2:
                    st.error("Not enough samples after preprocessing. Please check your data and missing value strategy.")
                    st.session_state.analysis_completed = False
                    st.stop()
                if len(np.unique(st.session_state.y_processed)) < 2:
                    st.error("The target variable must have at least two unique classes for classification. Please check your data or target selection.")
                    st.session_state.analysis_completed = False
                    st.stop()
                
                if st.session_state.label_encoder:
                    class_labels = list(st.session_state.label_encoder.classes_)
                else:
                    class_labels = [str(x) for x in np.unique(st.session_state.y_processed)]
                st.write(f"Detected classes in target variable: {', '.join(class_labels)}")

                st.header("4️⃣ Train-Test Split")
                st.markdown("""
                Divide your data into training and testing sets. The model learns from the training set
                and is evaluated on the unseen test set.
                """)
                test_size = st.slider("Select Test Data Size (%)", 10, 50, 20) / 100
                # Added the suggested tip line here
                st.markdown("💡 *A common split is 70% for training and 30% for testing (set slider to 30%).*")

                X_train, X_test, y_train, y_test = train_test_split(
                    st.session_state.X_processed, st.session_state.y_processed,
                    test_size=test_size, random_state=42, stratify=st.session_state.y_processed
                )
                st.write(f"Training set size: {X_train.shape[0]} samples")
                st.write(f"Test set size: {X_test.shape[0]} samples")

                st.header("5️⃣ Classification Model Evaluation")
                st.markdown("""
                We will now evaluate common classification models and perform a lightweight hyperparameter search
                to find the best performing model.
                """)

                models_and_params = {
                    "Logistic Regression": {
                        'model': LogisticRegression(random_state=42, solver='liblinear', max_iter=1000),
                        'params': {'C': [0.1, 1.0, 10.0]}
                    },
                    "Random Forest Classifier": {
                        'model': RandomForestClassifier(random_state=42),
                        'params': {'n_estimators': [50, 100, 150], 'max_depth': [5, 10, None]}
                    },
                    "Gradient Boosting Classifier": {
                        'model': GradientBoostingClassifier(random_state=42),
                        'params': {'n_estimators': [50, 100], 'learning_rate': [0.1, 0.2], 'max_depth': [3, 5]}
                    },
                }

                st.session_state.evaluation_results_with_best_params = {}
                
                if len(np.unique(y_train)) == 2:
                    scoring_metric = 'roc_auc'
                else:
                    scoring_metric = 'f1_weighted'

                st.write(f"Models will be evaluated using **{scoring_metric.replace('_', ' ').title()}** as the scoring metric.")

                evaluation_display_data = []

                with st.spinner("Performing automated model evaluation and hyperparameter search..."):
                    for model_name, config in models_and_params.items():
                        try:
                            model = config['model']
                            params = config['params']

                            grid_search = GridSearchCV(model, params, cv=3, scoring=scoring_metric, n_jobs=-1, verbose=0)
                            grid_search.fit(X_train, y_train)

                            best_model = grid_search.best_estimator_
                            best_params = grid_search.best_params_
                            best_score = grid_search.best_score_

                            y_pred_test = best_model.predict(X_test)
                            
                            accuracy_test = accuracy_score(y_test, y_pred_test)
                            precision_test = precision_score(y_test, y_pred_test, average='weighted', zero_division=0)
                            recall_test = recall_score(y_test, y_pred_test, average='weighted', zero_division=0)
                            f1_test = f1_score(y_test, y_pred_test, average='weighted', zero_division=0)
                            
                            roc_auc_test = np.nan
                            if len(np.unique(y_test)) == 2 and hasattr(best_model, "predict_proba"):
                                y_proba_test = best_model.predict_proba(X_test)[:, 1]
                                roc_auc_test = roc_auc_score(y_test, y_proba_test)

                            st.session_state.evaluation_results_with_best_params[model_name] = {
                                'best_model': best_model,
                                'best_params': best_params,
                                'best_cv_score': best_score,
                                'test_accuracy': accuracy_test,
                                'test_precision': precision_test,
                                'test_recall': recall_test,
                                'test_f1': f1_test,
                                'test_roc_auc': roc_auc_test,
                                'scoring_metric_used': scoring_metric
                            }
                            
                            evaluation_display_data.append({
                                "Model": model_name,
                                f"Best CV {scoring_metric.replace('_', ' ').title()}": best_score,
                                "Test Accuracy": accuracy_test,
                                "Test Precision": precision_test,
                                "Test Recall": recall_test,
                                "Test F1-Score": f1_test,
                                "Test ROC AUC": roc_auc_test if not np.isnan(roc_auc_test) else "N/A",
                                "Best Params": best_params
                            })

                        except Exception as e:
                            st.error(f"Error evaluating {model_name}: {e}")
                            st.session_state.evaluation_results_with_best_params[model_name] = {
                                'best_model': None, 'best_params': {}, 'best_cv_score': np.nan,
                                'test_accuracy': np.nan, 'test_precision': np.nan, 'test_recall': np.nan,
                                'test_f1': np.nan, 'test_roc_auc': np.nan, 'scoring_metric_used': scoring_metric
                            }
                            evaluation_display_data.append({
                                "Model": model_name,
                                f"Best CV {scoring_metric.replace('_', ' ').title()}": "Error",
                                "Test Accuracy": "Error", "Test Precision": "Error", "Test Recall": "Error",
                                "Test F1-Score": "Error", "Test ROC AUC": "Error", "Best Params": "Error"
                            })

                if evaluation_display_data:
                    results_df = pd.DataFrame(evaluation_display_data).set_index("Model")
                    st.dataframe(results_df.round(4))

                    st.markdown("""
                    **How to Interpret Metrics:**
                    - **Best CV Score:** Performance on training data via cross-validation (higher is better).
                    - **Test Metrics:** Performance on unseen test data (reflects generalization).
                    - **Accuracy:** Proportion of correct predictions.
                    - **Precision:** Of all positive predictions, how many were actually correct.
                    - **Recall:** Of all actual positives, how many were correctly identified.
                    - **F1-Score:** Harmonic mean of Precision and Recall, balancing both.
                    - **ROC AUC:** Measures the ability of the model to distinguish between classes (higher is better, 0.5 is random, 1.0 is perfect).
                    """)
                else:
                    st.info("No models could be evaluated. Please check your data and selections.")
                
                st.header("6️⃣ Model Recommendation")
                st.markdown("""
                Based on the automated evaluation, here is the recommended model and its best parameters.
                """)
                
                if st.session_state.evaluation_results_with_best_params:
                    valid_results = {k: v for k, v in st.session_state.evaluation_results_with_best_params.items() if not np.isnan(v['best_cv_score'])}
                    
                    if valid_results:
                        best_model_name = max(valid_results, key=lambda k: valid_results[k]['best_cv_score'])
                        best_model_info = valid_results[best_model_name]

                        st.session_state.recommended_model_info = {
                            'name': best_model_name,
                            'best_params': best_model_info['best_params'],
                            'score': best_model_info['best_cv_score'],
                            'metric': best_model_info['scoring_metric_used'].replace('_', ' ').title()
                        }

                        st.success(f"**Recommended Model:** {st.session_state.recommended_model_info['name']}")
                        st.write(f"**Best Parameters Found:** {st.session_state.recommended_model_info['best_params']}")
                        st.write(f"**Achieved Best CV {st.session_state.recommended_model_info['metric']}:** {format_metric(st.session_state.recommended_model_info['score'])}")
                        st.markdown("""
                        These parameters provide a strong starting point for the final model.
                        The sliders in the next section will be pre-filled with these values.
                        """)
                    else:
                        st.warning("No valid models could be recommended based on evaluation results.")
                        st.session_state.recommended_model_info = None
                else:
                    st.info("Run preprocessing and model evaluation to get a recommendation.")


                st.header("7️⃣ Choose Final Model for Classification")
                st.markdown("""
                Select your preferred model and adjust its hyperparameters for the final analysis.
                The sliders will suggest ranges around the recommended parameters.
                """)

                default_model_selection = "Logistic Regression"
                if st.session_state.recommended_model_info:
                    default_model_selection = st.session_state.recommended_model_info['name']

                chosen_classifier_name = st.selectbox(
                    "Select Classification Algorithm",
                    list(models_and_params.keys()),
                    index=list(models_and_params.keys()).index(default_model_selection) if default_model_selection in models_and_params.keys() else 0
                )

                final_model = None
                model_params = {}
                
                current_recommended_params = st.session_state.evaluation_results_with_best_params.get(chosen_classifier_name, {}).get('best_params', {})

                if chosen_classifier_name == "Logistic Regression":
                    default_c = current_recommended_params.get('C', 1.0)
                    min_c = max(0.01, default_c * 0.5)
                    max_c = default_c * 2.0
                    C_val = st.slider("Regularization Strength (C)", float(min_c), float(max_c), float(default_c), 0.01)
                    final_model = LogisticRegression(random_state=42, C=C_val, solver='liblinear', max_iter=1000)
                    model_params = {'C': C_val}

                elif chosen_classifier_name == "Random Forest Classifier":
                    default_n_estimators = current_recommended_params.get('n_estimators', 100)
                    default_max_depth = current_recommended_params.get('max_depth', 10)

                    min_n_estimators = max(10, default_n_estimators - 50)
                    max_n_estimators = default_n_estimators + 100
                    n_estimators = st.slider("Number of Estimators", min_n_estimators, max_n_estimators, default_n_estimators, 10)

                    min_max_depth = max(1, default_max_depth - 3 if default_max_depth is not None else 5)
                    max_max_depth = default_max_depth + 3 if default_max_depth is not None else 15
                    max_depth = st.slider("Max Depth (0 for unlimited)", min_max_depth, max_max_depth, default_max_depth if default_max_depth is not None else 0)
                    
                    final_model = RandomForestClassifier(random_state=42, n_estimators=n_estimators, max_depth=max_depth if max_depth > 0 else None)
                    model_params = {'n_estimators': n_estimators, 'max_depth': max_depth if max_depth > 0 else 'None'}

                elif chosen_classifier_name == "Gradient Boosting Classifier":
                    default_n_estimators_gb = current_recommended_params.get('n_estimators', 100)
                    default_learning_rate_gb = current_recommended_params.get('learning_rate', 0.1)
                    default_max_depth_gb = current_recommended_params.get('max_depth', 3)

                    min_n_estimators_gb = max(10, default_n_estimators_gb - 50)
                    max_n_estimators_gb = default_n_estimators_gb + 100
                    n_estimators_gb = st.slider("Number of Estimators (GB)", min_n_estimators_gb, max_n_estimators_gb, default_n_estimators_gb, 10)

                    min_learning_rate_gb = max(0.01, default_learning_rate_gb * 0.5)
                    max_learning_rate_gb = default_learning_rate_gb * 2.0
                    learning_rate_gb = st.slider("Learning Rate (GB)", float(min_learning_rate_gb), float(max_learning_rate_gb), float(default_learning_rate_gb), 0.01)

                    min_max_depth_gb = max(1, default_max_depth_gb - 2)
                    max_max_depth_gb = default_max_depth_gb + 2
                    max_depth_gb = st.slider("Max Depth (GB)", min_max_depth_gb, max_max_depth_gb, default_max_depth_gb)
                    
                    final_model = GradientBoostingClassifier(random_state=42, n_estimators=n_estimators_gb, learning_rate=learning_rate_gb, max_depth=max_depth_gb)
                    model_params = {'n_estimators': n_estimators_gb, 'learning_rate': learning_rate_gb, 'max_depth': max_depth_gb}

                st.markdown("""
                When you're ready, click the button below to train the final model and view results.
                """)

                if st.button("🚀 Run Classification"):
                    st.header("8️⃣ Classification Results")
                    with st.spinner(f"Training {chosen_classifier_name} and generating results..."):
                        fig_cm, fig_roc, fig_fi = None, None, None
                        try:
                            final_model.fit(X_train, y_train)
                            y_pred = final_model.predict(X_test)

                            final_accuracy = accuracy_score(y_test, y_pred)
                            final_precision = precision_score(y_test, y_pred, average='weighted', zero_division=0)
                            final_recall = recall_score(y_test, y_pred, average='weighted', zero_division=0)
                            final_f1 = f1_score(y_test, y_pred, average='weighted', zero_division=0)

                            final_roc_auc = np.nan
                            y_proba = None
                            if len(np.unique(y_test)) == 2 and hasattr(final_model, "predict_proba"):
                                y_proba = final_model.predict_proba(X_test)[:, 1]
                                final_roc_auc = roc_auc_score(y_test, y_proba)
                                st.success(f"✅ {chosen_classifier_name} training complete!")
                            elif len(np.unique(y_test)) > 2:
                                st.warning("ROC AUC is primarily for binary classification and not displayed for multi-class problems.")
                            else:
                                st.warning("ROC AUC not applicable for single-class target or model without probabilities.")


                            st.subheader("Performance Metrics on Test Set")
                            col_metrics1, col_metrics2, col_metrics3 = st.columns(3)
                            with col_metrics1:
                                st.metric("Accuracy", format_metric(final_accuracy))
                                st.metric("Precision", format_metric(final_precision))
                            with col_metrics2:
                                st.metric("Recall", format_metric(final_recall))
                                st.metric("F1-Score", format_metric(final_f1))
                            with col_metrics3:
                                if not np.isnan(final_roc_auc):
                                    st.metric("ROC AUC", format_metric(final_roc_auc))
                                else:
                                    st.info("ROC AUC not applicable.")

                            st.subheader("Confusion Matrix")
                            cm = confusion_matrix(y_test, y_pred)
                            fig_cm, ax_cm = plt.subplots(figsize=(8, 6))
                            
                            if st.session_state.label_encoder:
                                cm_labels = st.session_state.label_encoder.classes_
                            else:
                                cm_labels = [str(x) for x in np.unique(y_test)]

                            sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', ax=ax_cm,
                                        xticklabels=cm_labels,
                                        yticklabels=cm_labels)
                            ax_cm.set_xlabel('Predicted Label')
                            ax_cm.set_ylabel('True Label')
                            ax_cm.set_title('Confusion Matrix')
                            st.pyplot(fig_cm)
                            plt.close(fig_cm)

                            st.subheader("Classification Report")
                            if st.session_state.label_encoder:
                                y_test_labels = st.session_state.label_encoder.inverse_transform(y_test)
                                y_pred_labels = st.session_state.label_encoder.inverse_transform(y_pred)
                                target_names_report = st.session_state.label_encoder.classes_
                            else:
                                y_test_labels = y_test
                                y_pred_labels = y_pred
                                target_names_report = [str(x) for x in np.unique(y_test)]

                            class_report_dict_for_report = classification_report(y_test_labels, y_pred_labels, target_names=target_names_report, output_dict=True, zero_division=0)
                            st.text(classification_report(y_test_labels, y_pred_labels, target_names=target_names_report, zero_division=0))


                            st.subheader("Feature Importance")
                            feature_importances = None
                            if hasattr(final_model, 'feature_importances_'):
                                feature_importances = final_model.feature_importances_
                            elif hasattr(final_model, 'coef_'):
                                if final_model.coef_.ndim > 1:
                                    feature_importances = np.sum(np.abs(final_model.coef_), axis=0)
                                else:
                                    feature_importances = np.abs(final_model.coef_)
                                    
                            if feature_importances is not None and len(st.session_state.feature_names) == len(feature_importances):
                                importance_df = pd.DataFrame({
                                    'Feature': st.session_state.feature_names,
                                    'Importance': feature_importances
                                }).sort_values(by='Importance', ascending=False)

                                fig_fi, ax_fi = plt.subplots(figsize=(10, min(10, max(5, len(importance_df) * 0.4))))
                                sns.barplot(x='Importance', y='Feature', data=importance_df.head(15), ax=ax_fi)
                                ax_fi.set_title('Top Feature Importances')
                                plt.tight_layout()
                                st.pyplot(fig_fi)
                                plt.close(fig_fi)
                            else:
                                st.info("Feature importance is not available for this model or could not be calculated.")

                            st.session_state.analysis_completed = True

                            st.subheader("9️⃣ Download Results")
                            st.info("You can download the data with predictions or a comprehensive report.")

                            predictions_on_test_set_df = pd.DataFrame({
                                'True_Target_Encoded': y_test,
                                'Predicted_Target_Encoded': y_pred
                            }, index=X_test.index)

                            if st.session_state.label_encoder:
                                predictions_on_test_set_df['True_Target'] = st.session_state.label_encoder.inverse_transform(predictions_on_test_set_df['True_Target_Encoded'])
                                predictions_on_test_set_df['Predicted_Target'] = st.session_state.label_encoder.inverse_transform(predictions_on_test_set_df['Predicted_Target_Encoded'])
                            else:
                                predictions_on_test_set_df['True_Target'] = predictions_on_test_set_df['True_Target_Encoded']
                                predictions_on_test_set_df['Predicted_Target'] = predictions_on_test_set_df['Predicted_Target_Encoded']

                            df_with_predictions = st.session_state.original_df_with_target.copy()
                            df_with_predictions = df_with_predictions.merge(predictions_on_test_set_df[['True_Target', 'Predicted_Target']], left_index=True, right_index=True, how='left')
                            
                            df_with_predictions['True_Target'].fillna('N/A (Train/Dropped)', inplace=True)
                            df_with_predictions['Predicted_Target'].fillna('N/A (Train/Dropped)', inplace=True)


                            csv_buffer = io.StringIO()
                            df_with_predictions.to_csv(csv_buffer, index=False)
                            st.download_button(
                                label="📥 Download Data with Predictions (CSV)",
                                data=csv_buffer.getvalue(),
                                file_name=f"Classification_Predictions_{pd.to_datetime('today').strftime('%Y%m%d_%H%M%S')}.csv",
                                mime="text/csv",
                            )

                            cm_plot_bytes = io.BytesIO()
                            fig_cm.savefig(cm_plot_bytes, format='png', bbox_inches='tight')
                            cm_plot_bytes.seek(0)

                            roc_plot_bytes = io.BytesIO()
                            if len(np.unique(y_test)) == 2 and hasattr(final_model, "predict_proba") and y_proba is not None:
                                fpr, tpr, _ = roc_curve(y_test, y_proba)
                                roc_auc_val = auc(fpr, tpr)
                                fig_roc, ax_roc = plt.subplots(figsize=(8, 6))
                                ax_roc.plot(fpr, tpr, color='darkorange', lw=2, label=f'ROC curve (area = {roc_auc_val:.2f})')
                                ax_roc.plot([0, 1], [0, 1], color='navy', lw=2, linestyle='--')
                                ax_roc.set_xlabel('False Positive Rate')
                                ax_roc.set_ylabel('True Positive Rate')
                                ax_roc.set_title('Receiver Operating Characteristic (ROC) Curve')
                                ax_roc.legend(loc="lower right")
                                plt.close(fig_roc)
                                fig_roc.savefig(roc_plot_bytes, format='png', bbox_inches='tight')
                                roc_plot_bytes.seek(0)
                            else:
                                roc_plot_bytes = None

                            fi_plot_bytes = io.BytesIO()
                            if fig_fi:
                                fig_fi.savefig(fi_plot_bytes, format='png', bbox_inches='tight')
                                fi_plot_bytes.seek(0)
                            else:
                                fi_plot_bytes = None

                            report_bytes_io = io.BytesIO()
                            document = Document()
                            create_report(
                                document,
                                chosen_classifier_name,
                                model_params,
                                {
                                    'accuracy': final_accuracy,
                                    'precision': final_precision,
                                    'recall': final_recall,
                                    'f1_score': final_f1,
                                    'roc_auc': final_roc_auc,
                                    'feature_importances': feature_importances
                                },
                                df_with_predictions.head(5),
                                cm_plot_bytes.getvalue(),
                                roc_plot_bytes.getvalue() if roc_plot_bytes else None,
                                fi_plot_bytes.getvalue() if fi_plot_bytes else None,
                                class_report_dict_for_report,
                                st.session_state.target_name,
                                st.session_state.feature_names,
                                class_labels,
                                st.session_state.recommended_model_info
                            )
                            document.save(report_bytes_io)
                            report_bytes = report_bytes_io.getvalue()
                            report_bytes_io.close()

                            st.download_button(
                                label="📥 Download Comprehensive Report (.docx)",
                                data=report_bytes,
                                file_name=f"ML_Classification_Report_{pd.to_datetime('today').strftime('%Y%m%d_%H%M%S')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )

                        except Exception as e:
                            st.error(f"An error occurred during model training or result generation: {e}")
                            st.session_state.analysis_completed = False

            else:
                st.info("Please complete data preprocessing to proceed with model training.")
        else:
            st.info("Click '🚀 Run Classification' to see results.")

    if st.session_state.analysis_completed:
        st.header("🎯 Analysis Complete")
        st.markdown("You can either re-run classification with different parameters on the current dataset, or clear everything to start fresh with new data.")
        col_reset1, col_reset2 = st.columns(2)
        with col_reset1:
            if st.button("🔄 Rerun Classification with New Parameters"):
                st.session_state.analysis_completed = False
                st.rerun()
        with col_reset2:
            if st.button("🗑️ Clear All Data & Start Fresh"):
                auth_status = st.session_state.get('authenticated', False)
                current_file_uploader_key = st.session_state.get('file_uploader_key', 0)
                for key in list(st.session_state.keys()):
                    if key != 'authenticated':
                        del st.session_state[key]
                st.session_state.authenticated = auth_status
                st.session_state.file_uploader_key = current_file_uploader_key + 1
                st.rerun()
    else:
        if df is None:
            st.info("Please upload a data file (.csv or .xlsx) at the top of the page to begin the analysis.")

    st.markdown("---")
    st.markdown("<p style='text-align: center; color: grey;'>© Copyright SSO Consultants</p>", unsafe_allow_html=True)


if not st.session_state.authenticated:
    login_page()
else:
    main_app()
